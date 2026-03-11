import asyncio
import os
import logging
import re
from datetime import datetime
from typing import List, Dict, Any

# Библиотеки для Telegram
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup

# Библиотеки для работы с документами
import docx
import PyPDF2

class DocumentLoader:
    """Загружает текст и параграфы из .docx и .pdf файлов."""

    @staticmethod
    def _normalize_text(s: str) -> str:
        if s is None:
            return ""
        s = s.replace("\u00A0", " ").replace("\u200B", "").replace("\uFEFF", "")
        s = s.replace('\n', ' ').replace('\r', ' ').replace('\t', ' ')
        s = re.sub(r"[ \t\v\f\u00A0]+", " ", s)
        return s.strip()

    @staticmethod
    def load_docx_text_and_paragraphs(path: str, dedupe: bool = True, preserve_empty: bool = False):
        doc = docx.Document(path)
        paragraphs = []
        seen = set()

        def add_para(text):
            if text is None: text = ""
            t_norm = DocumentLoader._normalize_text(text) if text else ""
            if t_norm == "" and not preserve_empty: return

            if dedupe:
                if t_norm:
                    if t_norm not in seen:
                        paragraphs.append(t_norm)
                        seen.add(t_norm)
                elif preserve_empty:
                    paragraphs.append(t_norm)
            else:
                if t_norm == "" and not preserve_empty: return
                paragraphs.append(t_norm)

        for p in doc.paragraphs:
            text = "".join(run.text for run in p.runs)
            add_para(text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text = "".join(run.text for run in p.runs)
                        add_para(text)

        full_text = "\n".join(paragraphs)
        return full_text, paragraphs

    @staticmethod
    def load_pdf_text_and_paragraphs(path: str, dedupe: bool = True, preserve_empty: bool = False):
        text_lines = []
        seen = set()

        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text() or ""
                for ln in page_text.splitlines():
                    ln_norm = DocumentLoader._normalize_text(ln)
                    if ln_norm == "" and not preserve_empty: continue
                    if dedupe:
                        if ln_norm:
                            if ln_norm not in seen:
                                text_lines.append(ln_norm)
                                seen.add(ln_norm)
                        elif preserve_empty:
                            text_lines.append(ln_norm)
                    else:
                        if ln_norm == "" and not preserve_empty: continue
                        text_lines.append(ln_norm)

        full_text = "\n".join(text_lines)
        return full_text, text_lines

    @staticmethod
    def get_paragraphs(path: str):
        lower = path.lower()
        if lower.endswith(".docx"):
            _, paras = DocumentLoader.load_docx_text_and_paragraphs(path, dedupe=False, preserve_empty=True)
            return paras
        elif lower.endswith(".pdf"):
            _, paras = DocumentLoader.load_pdf_text_and_paragraphs(path, dedupe=False, preserve_empty=True)
            return paras
        else:
            raise ValueError("Поддерживаются только .docx и .pdf")


class Template:
    """Хранит список placeholders и их anchors."""

    def __init__(self, placeholders=None, source_path=None):
        self.placeholders = placeholders or []
        self.source_path = source_path

    def get_placeholders(self):
        return self.placeholders

    @staticmethod
    def _normalize_type(raw_type: str) -> str:
        t = raw_type.strip().lower()
        if t in ("int", "integer", "num", "number", "float"):
            return "number"
        if t in ("str", "string", "text"):
            return "string"
        if t in ("date", "dt"):
            return "date"
        return t

    @staticmethod
    def extract_placeholders_from_paragraphs(paragraphs: list) -> list:
        placeholders = []
        inline_pattern = re.compile(
            r"\[\[\s*([^:\]\n]+?)\s*:\s*([^,:]\s*[^,\]\n]+?)"
            r"(?:\s*:\s*([^:\]\n]+?)\s*:\s*([^,\]\n]+?))?"
            r"(?:\s*,\s*(optional))?\s*\]\]",
            flags=re.IGNORECASE,
        )

        skip_patterns = [
            "утверждаю", "задание", "введение", "заключение",
            "список использованных источников", "примерный календарный график",
            "подпись обучающегося",
        ]

        for idx, para in enumerate(paragraphs):
            if not para.strip():
                continue

            for m in inline_pattern.finditer(para):
                raw_name = m.group(1).strip()
                raw_type = m.group(2).strip()
                raw_group_name = m.group(3)
                raw_group_condition = m.group(4)
                optional_flag = m.group(5)

                is_group_defined = raw_group_name is not None
                if is_group_defined:
                    raw_group_name = raw_group_name.strip() if raw_group_name else ""
                    raw_group_condition = raw_group_condition.strip() if raw_group_condition else ""
                else:
                    raw_group_name = ""
                    raw_group_condition = ""

                optional_flag = bool(optional_flag)

                # anchor_before
                left_part = para[: m.start()].strip()
                if left_part:
                    anchor_before = left_part
                else:
                    anchor_before = ""
                    for j in range(idx - 1, -1, -1):
                        prev_para = paragraphs[j].strip()
                        if (
                                prev_para and not inline_pattern.search(prev_para)
                                and not any(sp in prev_para.lower() for sp in skip_patterns)
                        ):
                            anchor_before = prev_para
                            break

                # anchor_after
                right_part = para[m.end():].strip()
                if right_part:
                    anchor_after = right_part
                else:
                    anchor_after = ""
                    max_template_distance = 6
                    forbidden_after = [
                        "(подпись)", "(инициалы", "фамилия", "подпись", "инициалы",
                        "(инициалы, фамилия)", "подпись обучающегося", "(подпись обучающегося)",
                    ]
                    for j in range(idx + 1, min(len(paragraphs), idx + 1 + max_template_distance)):
                        next_para = paragraphs[j].strip()
                        if not next_para: continue
                        next_lower = next_para.lower()
                        if any(f in next_lower for f in forbidden_after): continue
                        if (
                                not inline_pattern.search(next_para)
                                and not any(sp in next_lower for sp in skip_patterns)
                        ):
                            anchor_after = next_para
                            break

                placeholders.append(
                    {
                        "name": raw_name,
                        "type": Template._normalize_type(raw_type),
                        "optional": optional_flag,
                        "group_name": raw_group_name,
                        "group_condition": raw_group_condition,
                        "anchor_before": anchor_before,
                        "anchor_after": anchor_after,
                        "source_paragraph": para,
                        "para_index": idx,
                    }
                )

        seen = set()
        unique = []
        for p in placeholders:
            key = (
                p["name"].lower(),
                re.sub(r"\s+", " ", p["anchor_before"].strip()).lower() if p["anchor_before"] else "",
                re.sub(r"\s+", " ", p["anchor_after"].strip()).lower() if p["anchor_after"] else "",
                p["group_name"].lower(), p["group_condition"].lower(),
            )
            if key not in seen:
                seen.add(key)
                unique.append(p)

        for i in range(len(unique) - 1):
            unique[i]["next_is_placeholder"] = unique[i + 1]["para_index"] - unique[i]["para_index"] <= 1

        return unique

    @classmethod
    def load_from_file(cls, path: str):
        lower = path.lower()
        if lower.endswith(".docx"):
            _, paragraphs = DocumentLoader.load_docx_text_and_paragraphs(path, dedupe=True, preserve_empty=False)
        elif lower.endswith(".pdf"):
            _, paragraphs = DocumentLoader.load_pdf_text_and_paragraphs(path, dedupe=True, preserve_empty=False)
        else:
            raise ValueError("Поддерживаются только .docx и .pdf")

        placeholders = cls.extract_placeholders_from_paragraphs(paragraphs)
        if not placeholders:
            raise ValueError("В шаблоне не найдено ни одного заполнителя [[...]]")
        return cls(placeholders=placeholders, source_path=path)


class DocumentChecker:
    """Проверяет документ по anchors и выполняет групповые проверки."""

    def __init__(self, template: Template):
        self.template = template
        self._placeholder_pattern = re.compile(r"\[\[.*?\]\]")

    @staticmethod
    def _extract_first_number(value: str) -> str | None:
        if not value:
            return None
        m = re.search(r"([+-]?\s*\d+([.,]\d+)?)", value)
        if m:
            return m.group(1).strip()
        return None

    @staticmethod
    def _validate_type(value: str, expected_type: str) -> bool:
        if not value: return False
        v = value.strip()
        if expected_type == "string": return bool(re.search(r"[A-Za-zА-Яа-яЁё]", v))
        if expected_type == "number": return bool(
            re.fullmatch(r"[+-]?\s*\d+([.,]\d+)?", v.replace(' ', '')))
        if expected_type == "date": return bool(
            re.fullmatch(r"\d{1,2}\.\d{1,2}\.\d{4}", v) or re.search(r"\d{1,2}\s+[А-Яа-яёЁ]+\.?\s+\d{4}", v))
        return True

    @staticmethod
    def _is_anchor_like(value: str, anchors: list) -> bool:
        if not value: return False
        v = re.sub(r"\s+", " ", value).strip().lower()
        for a in anchors:
            if not a: continue
            a_norm = re.sub(r"\s+", " ", a).strip().lower()
            if v == a_norm: return True
        return False

    def _find_value_using_anchors(
            self, anchor_before, anchor_after, doc_paragraphs, start_index=0, expected_type=None,
            next_is_placeholder=False
    ):
        stop_words_list = [
            "введение", "заключение", "список использованных источников",
            "примерный календарный график", "приложение", "руководитель курсового проекта",
            "куратор", "проверяющий", "обучающемуся", "задание",
        ]

        def find_positions(anchor):
            if not anchor: return []
            a_norm = re.sub(r'["\s]+', ' ', anchor.strip()).lower()
            pos = []
            for i in range(start_index, len(doc_paragraphs)):
                para = doc_paragraphs[i]
                if para is None: continue
                para_norm_cells = re.sub(r"[,;]+", "|", para.strip()).lower()
                para_norm_text = re.sub(r"\s+", " ", para.strip()).lower()
                if a_norm in para_norm_text or a_norm in para_norm_cells.split('|'):
                    pos.append(i)
            return pos

        pos_before = find_positions(anchor_before)
        pos_after = find_positions(anchor_after)

        def candidate_ok(val, anchors):
            if not val: return False
            v = val.strip()
            if v == "": return False
            if self._is_anchor_like(v, anchors): return False
            if re.match(r"^\d+\.", v): return False
            if any(sw in v.lower() for sw in stop_words_list): return False
            if re.search(r"подпись|иниц|инициалы|фамил", v.lower()): return False
            if expected_type and expected_type != 'number': return self._validate_type(v, expected_type)
            return True

        if expected_type == 'number':
            max_forward_search = 10
            start_pos = start_index
            if pos_before:
                start_pos = pos_before[0] + 1

            for k in range(start_pos, min(len(doc_paragraphs), start_pos + max_forward_search)):
                cand_raw = doc_paragraphs[k]
                if cand_raw is None or cand_raw.strip() == "": continue
                cand = cand_raw.strip()
                cand_lower = cand.lower()

                if any(sw in cand_lower for sw in stop_words_list) or re.search(
                        r"^\(?\s*(подпись|иниц|инициалы|фамил)", cand_lower):
                    break
                if self._placeholder_pattern.search(cand):
                    break

                extracted_number = self._extract_first_number(cand)
                if extracted_number:
                    if not self._is_anchor_like(extracted_number, [anchor_before, anchor_after]):
                        return True, extracted_number, k

                if not extracted_number and len(cand) > 30:
                    break
                continue
            return False, None, -1

        if pos_before and pos_after:
            best = None
            best_dist = None
            for b in pos_before:
                for a in pos_after:
                    if b >= a: continue
                    dist = a - b
                    if best_dist is None or dist < best_dist:
                        best_dist = dist
                        best = (b, a)
            if best:
                b, a = best
                max_doc_distance = 8
                if a - b <= max_doc_distance:
                    if (a - b == 1) or next_is_placeholder:
                        pass
                    else:
                        for k in range(b + 1, a):
                            mid = doc_paragraphs[k].strip()
                            if candidate_ok(mid, [anchor_before, anchor_after]):
                                return True, mid, k

        if pos_before:
            for b in pos_before:
                para_b = doc_paragraphs[b] or ""
                low_b = para_b.lower()
                ab = anchor_before.strip().lower()
                idx_b = low_b.find(ab)
                if idx_b != -1:
                    after_b = para_b[idx_b + len(ab):].strip()
                    if candidate_ok(after_b, [anchor_before, anchor_after]):
                        return True, after_b, b

                max_forward_search = 10
                for k in range(b + 1, min(len(doc_paragraphs), b + 1 + max_forward_search)):
                    cand_raw = doc_paragraphs[k]
                    if cand_raw is None or cand_raw.strip() == "": continue
                    cand = cand_raw.strip()
                    cand_lower = cand.lower()
                    if any(sw in cand_lower for sw in ["введение", "заключение", "список", "приложение"]) or re.search(
                            r"^\(?\s*(подпись|иниц|инициалы|фамил)", cand_lower):
                        break
                    if candidate_ok(cand, [anchor_before, anchor_after]):
                        return True, cand, k
                    break

        if pos_after:
            for a in pos_after:
                para_a = doc_paragraphs[a] or ""
                low_a = para_a.lower()
                aa = anchor_after.strip().lower()
                idx_a = low_a.find(aa)
                if idx_a != -1:
                    before_a = para_a[:idx_a].strip()
                    if candidate_ok(before_a, [anchor_before, anchor_after]):
                        return True, before_a, a
                k = a - 1
                if k >= 0:
                    cand_raw = doc_paragraphs[k]
                    if cand_raw and cand_raw.strip():
                        cand = cand_raw.strip()
                        if not (any(sw in cand.lower() for sw in stop_words_list) or re.search(
                                r"^\(?\s*(подпись|иниц|инициалы|фамил)", cand.lower())):
                            if candidate_ok(cand, [anchor_before, anchor_after]):
                                return True, cand, k
        return False, None, -1

    def _evaluate_group_condition(self, condition: str, sum_val: float, num_values: int) -> tuple:
        if not condition: return True, "Нет условия для проверки."
        condition = condition.strip().upper().replace(' ', '')
        m = re.match(r"(SUM|AVG)([<=>!]+)(\d+(\.\d+)?)", condition)
        if not m: return False, f"Условие '{condition}' не распознано."

        check_type, operator, target_str = m.group(1), m.group(2), m.group(3)
        target = float(target_str)
        value_to_check = sum_val
        check_name = "Сумма"
        if check_type == "AVG":
            if num_values == 0: return False, "Деление на ноль (нет значений)."
            value_to_check = sum_val / num_values
            check_name = "Среднее"

        tolerance = 0.001
        is_valid = False
        if operator == '=':
            is_valid = abs(value_to_check - target) < tolerance
        elif operator == '>=':
            is_valid = value_to_check >= target
        elif operator == '<=':
            is_valid = value_to_check <= target
        elif operator == '>':
            is_valid = value_to_check > target
        elif operator == '<':
            is_valid = value_to_check < target
        elif operator == '!=':
            is_valid = abs(value_to_check - target) >= tolerance

        result_val_str = f"{value_to_check:.2f}"
        msg = f"{check_name} ({result_val_str}) соответствует условию {check_type}{operator}{target}." if is_valid else \
            f"{check_name} ({result_val_str}) не соответствует условию {check_type}{operator}{target}."
        return is_valid, msg

    def _check_groups(self, results: list) -> list:
        groups_to_check = {}
        for r in results:
            g_name = r.get("group_name")
            g_cond = r.get("group_condition", "").strip()
            if g_name and g_cond:
                key = (g_name, g_cond)
                if key not in groups_to_check:
                    groups_to_check[key] = {
                        "condition": g_cond, "valid_values": [], "total_sum": 0.0,
                        "all_ok": True, "missing_fields": []
                    }
                if r['status'] == 'ok' and r['expected_type'] == 'number':
                    try:
                        val_str = str(r['value']).replace(' ', '').replace(',', '.')
                        float_value = float(val_str)
                        groups_to_check[key]["valid_values"].append(float_value)
                        groups_to_check[key]["total_sum"] += float_value
                    except ValueError:
                        groups_to_check[key]["all_ok"] = False
                        groups_to_check[key]["missing_fields"].append(f"{r['field']} (не число)")
                else:
                    groups_to_check[key]["all_ok"] = False
                    groups_to_check[key]["missing_fields"].append(r['field'])

        group_report = []
        for (group_name, condition), data in groups_to_check.items():
            if data["all_ok"]:
                is_valid, message = self._evaluate_group_condition(condition, data["total_sum"],
                                                                   len(data["valid_values"]))
                status = "group_ok" if is_valid else "group_condition_invalid"
            else:
                status = "group_check_failed"
                missing = ", ".join(data['missing_fields'])
                message = f"Ошибки в полях: {missing}."

            group_report.append({
                "field": f"Группа: {group_name}", "status": status,
                "value": f"{data['total_sum']:.2f}", "group_name": group_name,
                "message": message
            })
        return group_report

    def check_document(self, doc_paragraphs: list) -> list:
        results = []
        cursor = 0
        max_len = len(doc_paragraphs)
        while cursor < max_len and not doc_paragraphs[cursor].strip(): cursor += 1

        for ph in self.template.get_placeholders():
            found, value, found_idx = self._find_value_using_anchors(
                ph.get("anchor_before", ""), ph.get("anchor_after", ""),
                doc_paragraphs, cursor, ph["type"], ph.get("next_is_placeholder", False)
            )

            if not found:
                status = "missing_optional" if ph["optional"] else "missing"
                results.append({
                    "field": ph["name"], "status": status, "optional": ph["optional"],
                    "group_name": ph["group_name"], "group_condition": ph["group_condition"]
                })
                if ph["type"] == 'number' and ph["name"].startswith("Этап_"):
                    cursor = max(cursor, cursor + 1)
                continue

            is_valid = self._validate_type(value, ph["type"]) if ph["type"] != 'number' else True
            results.append({
                "field": ph["name"], "value": value, "expected_type": ph["type"],
                "status": "ok" if is_valid else "invalid", "optional": ph["optional"],
                "group_name": ph["group_name"], "group_condition": ph["group_condition"]
            })
            cursor = max(cursor, found_idx + 1)
            while cursor < max_len and not doc_paragraphs[cursor].strip(): cursor += 1

        return results + self._check_groups(results)

    # -------------------------------------------------------------
    # ИЗМЕНЕННЫЙ МЕТОД: Добавлено отображение группы рядом с именем
    # -------------------------------------------------------------
    def generate_report(self, file_name: str, results: list) -> str:
        lines = [f"📄 Файл: {file_name}", ""]
        group_reports = []

        for r in results:
            # Отделяем отчеты о группах в отдельный список
            if r["field"].startswith("Группа:"):
                group_reports.append(r)
                continue

            # --- ИЗМЕНЕНИЕ: Формируем строку с названием группы ---
            group_info = ""
            if r.get("group_name"):
                group_info = f" <i>(Группа: {r['group_name']})</i>"
            # -----------------------------------------------------

            if r["status"] == "ok":
                lines.append(f"✅ <b>{r['field']}</b>{group_info}: {r['value']}")
            elif r["status"] == "invalid":
                lines.append(f"⚠️ <b>{r['field']}</b>{group_info}: '{r['value']}' (Тип не {r['expected_type']})")
            elif r["status"] == "missing_optional":
                lines.append(f"ℹ️ {r['field']}{group_info}: пропущено (необяз.)")
            elif r["status"] == "missing":
                lines.append(f"❌ <b>{r['field']}</b>{group_info}: Не найдено")

        if group_reports:
            lines.append("\n<b>Группы и формулы:</b>")
            for gr in group_reports:
                icon = "✅" if gr['status'] == 'group_ok' else "❌"
                if gr['status'] == 'group_check_failed': icon = "⚠️"
                lines.append(f"{icon} {gr['field']}: {gr['message']}")

        return "\n".join(lines)


class MultiAgentCheckSystem:
    def process(self, template_path: str, doc_path: str) -> str:
        try:
            tpl = Template.load_from_file(template_path)
            doc_paras = DocumentLoader.get_paragraphs(doc_path)

            checker = DocumentChecker(tpl)
            results = checker.check_document(doc_paras)
            return checker.generate_report(os.path.basename(doc_path), results)

        except Exception as e:
            logging.error(f"Error: {e}", exc_info=True)
            return f"Критическая ошибка: {str(e)}"


BOT_TOKEN = "8124707173:AAEUWIG6cU8ErdX_ItQZdbWNGD3JRLwjjNo"  

logging.basicConfig(level=logging.INFO)
bot = Bot(token=BOT_TOKEN)
dp = Dispatcher(storage=MemoryStorage())
system = MultiAgentCheckSystem()

TEMP_DIR = "temp_files"
os.makedirs(TEMP_DIR, exist_ok=True)


class Workflow(StatesGroup):
    waiting_for_template = State()
    waiting_for_document = State()


@dp.message(Command("start"))
async def cmd_start(message: types.Message, state: FSMContext):
    await message.answer(
        "👋 Привет! Я <b>САПСР (Fixed Edition)</b>.\n\n"
        "Пришлите <b>ШАБЛОН</b> в формате docx/pdf\n",
        parse_mode="HTML"
    )
    await state.set_state(Workflow.waiting_for_template)


@dp.message(Workflow.waiting_for_template, F.document)
async def process_template(message: types.Message, state: FSMContext):
    file_name = message.document.file_name
    if not (file_name.endswith('.docx') or file_name.endswith('.pdf')):
        await message.answer("❌ Поддерживаются только .docx и .pdf файлы.")
        return

    file = await bot.get_file(message.document.file_id)
    local_path = os.path.join(TEMP_DIR, f"tpl_{message.from_user.id}_{file_name}")
    await bot.download_file(file.file_path, local_path)

    await state.update_data(template_path=local_path)
    await message.answer(f"✅ Шаблон <b>{file_name}</b> загружен. \nЖду документ для проверки.", parse_mode="HTML")
    await state.set_state(Workflow.waiting_for_document)


@dp.message(Workflow.waiting_for_document, F.document)
async def process_document(message: types.Message, state: FSMContext):
    data = await state.get_data()
    template_path = data.get("template_path")
    if not template_path:
        await message.answer("⚠️ Шаблон потерян. Начните с /start")
        return

    msg = await message.answer("⏳ Обработка данных...")

    file_name = message.document.file_name
    file = await bot.get_file(message.document.file_id)
    doc_path = os.path.join(TEMP_DIR, f"doc_{message.from_user.id}_{file_name}")
    await bot.download_file(file.file_path, doc_path)

    # Запуск логики проверки в потоке (так как parsing может быть тяжелым)
    report = await asyncio.to_thread(system.process, template_path, doc_path)

    if len(report) > 4000:
        for x in range(0, len(report), 4000):
            await message.answer(report[x:x + 4000], parse_mode="HTML")
    else:
        await msg.edit_text(report, parse_mode="HTML")

    await message.answer("Можете прислать следующий документ или /start для смены шаблона.")


@dp.message(Command("cancel"))
async def cmd_cancel(message: types.Message, state: FSMContext):
    await state.clear()
    await message.answer("Сброс выполнен. Жмите /start")


async def main():
    await bot.delete_webhook(drop_pending_updates=True)
    await dp.start_polling(bot)


if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        pass
