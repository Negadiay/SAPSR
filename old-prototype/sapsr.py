# -*- coding: utf-8 -*-
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import docx
import PyPDF2
import os
import re
from datetime import datetime



class DocumentLoader:
    """Загружает текст и параграфы из .docx и .pdf файлов."""

    @staticmethod
    def _normalize_text(s: str) -> str:
        if s is None:
            return ""
        s = s.replace("\u00A0", " ").replace("\u200B", "").replace("\uFEFF", "")
        s = s.replace('\n', ' ').replace('\r', ' ').replace('\t', ' ')
        s = re.sub(r"[ \t\v\f\u00A0]+", " ", s)
        return s.strip()

    @staticmethod
    def load_docx_text_and_paragraphs(path: str, dedupe: bool = True, preserve_empty: bool = False):
        doc = docx.Document(path)
        paragraphs = []
        seen = set()

        def add_para(text):
            if text is None: text = ""
            t_norm = DocumentLoader._normalize_text(text) if text else ""
            if t_norm == "" and not preserve_empty: return

            if dedupe:
                if t_norm:
                    if t_norm not in seen:
                        paragraphs.append(t_norm)
                        seen.add(t_norm)
                elif preserve_empty:
                    paragraphs.append(t_norm)
            else:
                if t_norm == "" and not preserve_empty: return
                paragraphs.append(t_norm)

        for p in doc.paragraphs:
            text = "".join(run.text for run in p.runs)
            add_para(text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        text = "".join(run.text for run in p.runs)
                        add_para(text)

        full_text = "\n".join(paragraphs)
        return full_text, paragraphs

    @staticmethod
    def load_pdf_text_and_paragraphs(path: str, dedupe: bool = True, preserve_empty: bool = False):
        text_lines = []
        seen = set()

        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text() or ""
                for ln in page_text.splitlines():
                    ln_norm = DocumentLoader._normalize_text(ln)

                    if ln_norm == "" and not preserve_empty: continue

                    if dedupe:
                        if ln_norm:
                            if ln_norm not in seen:
                                text_lines.append(ln_norm)
                                seen.add(ln_norm)
                        elif preserve_empty:
                            text_lines.append(ln_norm)
                    else:
                        if ln_norm == "" and not preserve_empty: continue
                        text_lines.append(ln_norm)

        full_text = "\n".join(text_lines)
        return full_text, text_lines

    @staticmethod
    def get_paragraphs(path: str):
        lower = path.lower()
        if lower.endswith(".docx"):
            # Для docx используем preserve_empty=True, чтобы не потерять пустые абзацы, которые могут быть разделителями
            _, paras = DocumentLoader.load_docx_text_and_paragraphs(path, dedupe=False, preserve_empty=True)
            return paras
        elif lower.endswith(".pdf"):
            _, paras = DocumentLoader.load_pdf_text_and_paragraphs(path, dedupe=False, preserve_empty=True)
            return paras
        else:
            raise ValueError("Поддерживаются только .docx и .pdf")

class Template:
    """Хранит список placeholders и их anchors."""

    def __init__(self, placeholders=None, source_path=None):
        self.placeholders = placeholders or []
        self.source_path = source_path

    def get_placeholders(self):
        return self.placeholders

    @staticmethod
    def _normalize_type(raw_type: str) -> str:
        t = raw_type.strip().lower()
        if t in ("int", "integer", "num", "number", "float"):
            return "number"
        if t in ("str", "string", "text"):
            return "string"
        if t in ("date", "dt"):
            return "date"
        return t

    @staticmethod
    def extract_placeholders_from_paragraphs(paragraphs: list) -> list:
        """
        Находит placeholders. Поддерживает [[name:type(:group_name:group_condition)?(, optional)?]]
        """
        placeholders = []

        inline_pattern = re.compile(
            r"\[\[\s*([^:\]\n]+?)\s*:\s*([^,:]\s*[^,\]\n]+?)"
            r"(?:\s*:\s*([^:\]\n]+?)\s*:\s*([^,\]\n]+?))?"
            r"(?:\s*,\s*(optional))?\s*\]\]",
            flags=re.IGNORECASE,
        )

        skip_patterns = [
            "утверждаю", "задание", "введение", "заключение",
            "список использованных источников", "примерный календарный график",
            "подпись обучающегося",
        ]

        for idx, para in enumerate(paragraphs):
            if not para.strip():
                continue

            for m in inline_pattern.finditer(para):
                raw_name = m.group(1).strip()
                raw_type = m.group(2).strip()
                raw_group_name = m.group(3)
                raw_group_condition = m.group(4)
                optional_flag = m.group(5)

                is_group_defined = raw_group_name is not None

                if is_group_defined:
                    raw_group_name = raw_group_name.strip() if raw_group_name else ""
                    raw_group_condition = raw_group_condition.strip() if raw_group_condition else ""
                else:
                    raw_group_name = ""
                    raw_group_condition = ""

                optional_flag = bool(optional_flag)

                # ---------------- anchor_before ----------------
                left_part = para[: m.start()].strip()
                if left_part:
                    anchor_before = left_part
                else:
                    anchor_before = ""
                    for j in range(idx - 1, -1, -1):
                        prev_para = paragraphs[j].strip()
                        if (
                                prev_para and not inline_pattern.search(prev_para)
                                and not any(sp in prev_para.lower() for sp in skip_patterns)
                        ):
                            anchor_before = prev_para
                            break

                # ---------------- anchor_after ----------------
                right_part = para[m.end():].strip()
                if right_part:
                    anchor_after = right_part
                else:
                    anchor_after = ""
                    max_template_distance = 6
                    forbidden_after = [
                        "(подпись)", "(инициалы", "фамилия", "подпись", "инициалы",
                        "(инициалы, фамилия)", "подпись обучающегося", "(подпись обучающегося)",
                    ]
                    for j in range(idx + 1, min(len(paragraphs), idx + 1 + max_template_distance)):
                        next_para = paragraphs[j].strip()
                        if not next_para: continue
                        next_lower = next_para.lower()
                        if any(f in next_lower for f in forbidden_after): continue
                        if (
                                not inline_pattern.search(next_para)
                                and not any(sp in next_lower for sp in skip_patterns)
                        ):
                            anchor_after = next_para
                            break

                placeholders.append(
                    {
                        "name": raw_name,
                        "type": Template._normalize_type(raw_type),
                        "optional": optional_flag,
                        "group_name": raw_group_name,
                        "group_condition": raw_group_condition,
                        "anchor_before": anchor_before,
                        "anchor_after": anchor_after,
                        "source_paragraph": para,
                        "para_index": idx,
                    }
                )

        # удаление дубликатов
        seen = set()
        unique = []
        for p in placeholders:
            key = (
                p["name"].lower(),
                re.sub(r"\s+", " ", p["anchor_before"].strip()).lower() if p["anchor_before"] else "",
                re.sub(r"\s+", " ", p["anchor_after"].strip()).lower() if p["anchor_after"] else "",
                p["group_name"].lower(), p["group_condition"].lower(),
            )
            if key not in seen:
                seen.add(key)
                unique.append(p)
                
        for i in range(len(unique) - 1):
            unique[i]["next_is_placeholder"] = unique[i + 1]["para_index"] - unique[i]["para_index"] <= 1

        return unique

    @classmethod
    def load_from_file(cls, path: str):
        lower = path.lower()
        if lower.endswith(".docx"):
            _, paragraphs = DocumentLoader.load_docx_text_and_paragraphs(path, dedupe=True, preserve_empty=False)
        elif lower.endswith(".pdf"):
            _, paragraphs = DocumentLoader.load_pdf_text_and_paragraphs(path, dedupe=True, preserve_empty=False)
        else:
            raise ValueError("Поддерживаются только .docx и .pdf")

        placeholders = cls.extract_placeholders_from_paragraphs(paragraphs)
        if not placeholders:
            raise ValueError("В шаблоне не найдено ни одного заполнителя [[...]]")
        return cls(placeholders=placeholders, source_path=path)

class DocumentChecker:
    """Проверяет документ по anchors и выполняет групповые проверки."""

    def __init__(self, template: Template):
        self.template = template
        self._placeholder_pattern = re.compile(r"\[\[.*?\]\]")

    @staticmethod
    def _extract_first_number(value: str) -> str | None:
        """Извлекает первое число (целое или десятичное) из строки."""
        if not value:
            return None
        # Ищем число в формате [+-]?\s*\d+([.,]\d+)?
        m = re.search(r"([+-]?\s*\d+([.,]\d+)?)", value)
        if m:
            # Возвращаем найденную группу, очищенную от пробелов
            return m.group(1).strip()
        return None

    @staticmethod
    def _validate_type(value: str, expected_type: str) -> bool:
        if not value: return False
        v = value.strip()
        if expected_type == "string": return bool(re.search(r"[A-Za-zА-Яа-яЁё]", v))
        if expected_type == "number": return bool(
            re.fullmatch(r"[+-]?\s*\d+([.,]\d+)?", v.replace(' ', '')))
        if expected_type == "date": return bool(
            re.fullmatch(r"\d{1,2}\.\d{1,2}\.\d{4}", v) or re.search(r"\d{1,2}\s+[А-Яа-яёЁ]+\.?\s+\d{4}", v))
        return True

    @staticmethod
    def _is_anchor_like(value: str, anchors: list) -> bool:
        if not value: return False
        v = re.sub(r"\s+", " ", value).strip().lower()
        for a in anchors:
            if not a: continue
            a_norm = re.sub(r"\s+", " ", a).strip().lower()
            if v == a_norm: return True
        return False

    def _find_value_using_anchors(
            self, anchor_before, anchor_after, doc_paragraphs, start_index=0, expected_type=None,
            next_is_placeholder=False
    ):
        stop_words_list = [
            "введение", "заключение", "список использованных источников",
            "примерный календарный график", "приложение", "руководитель курсового проекта",
            "куратор", "проверяющий", "обучающемуся", "задание",
        ]

        def find_positions(anchor):
            if not anchor: return []
            a_norm = re.sub(r'["\s]+', ' ', anchor.strip()).lower()
            pos = []
            for i in range(start_index, len(doc_paragraphs)):
                para = doc_paragraphs[i]
                if para is None: continue

                para_norm_cells = re.sub(r"[,;]+", "|", para.strip()).lower()
                para_norm_text = re.sub(r"\s+", " ", para.strip()).lower()

                # Проверка:
                # 1. Точное вхождение в нормализованный текст
                # 2. Вхождение как отдельная "ячейка" (с учетом разделителей ",")
                if a_norm in para_norm_text or a_norm in para_norm_cells.split('|'):
                    pos.append(i)
            return pos

        pos_before = find_positions(anchor_before)
        pos_after = find_positions(anchor_after)

        def candidate_ok(val, anchors):
            if not val: return False
            v = val.strip()
            if v == "": return False
            # Проверка, не является ли значение самим якорем
            if self._is_anchor_like(v, anchors): return False
            # Проверка на список или заголовок
            if re.match(r"^\d+\.", v): return False
            # Проверка на общие стоп-слова
            if any(sw in v.lower() for sw in stop_words_list): return False
            if re.search(r"подпись|иниц|инициалы|фамил", v.lower()): return False
            # Проверка типа (только если не числовое поле)
            if expected_type and expected_type != 'number': return self._validate_type(v, expected_type)
            # Для числовых полей возвращаем True, т.к. извлечение числа происходит отдельно
            return True

            if expected_type == 'number':
                
                if pos_before:
                    idx = pos_before[0]
                    current_para = doc_paragraphs[idx] or ""
                    anchor_clean = anchor_before.strip().lower()
                    para_lower = current_para.lower()
                    find_idx = para_lower.find(anchor_clean)

                    if find_idx != -1:
                        text_after_anchor = current_para[find_idx + len(anchor_clean):]
                        extracted_number_inline = self._extract_first_number(text_after_anchor)

                        if extracted_number_inline:
                            return True, extracted_number_inline, idx

                max_forward_search = 10
                start_pos = start_index

                if pos_before:
                    start_pos = pos_before[0] + 1

                for k in range(start_pos, min(len(doc_paragraphs), start_pos + max_forward_search)):
                    cand_raw = doc_paragraphs[k]

                    if cand_raw is None or cand_raw.strip() == "":
                        continue

                    cand = cand_raw.strip()
                    cand_lower = cand.lower()

                    if any(sw in cand_lower for sw in stop_words_list) or re.search(
                            r"^\(?\s*(подпись|иниц|инициалы|фамил)", cand_lower):
                        break

                    if self._placeholder_pattern.search(cand):
                        break

                    extracted_number = self._extract_first_number(cand)

                    if extracted_number:
                        if not self._is_anchor_like(extracted_number, [anchor_before, anchor_after]):
                            return True, extracted_number, k

                    # Если не нашли число и абзац содержит много текста, предполагаем выход из таблицы
                    if not extracted_number and len(cand) > 30:
                        break

                    continue

                return False, None, -1


        # 1. Поиск по двум якорям
        if pos_before and pos_after:
            best = None
            best_dist = None
            for b in pos_before:
                for a in pos_after:
                    if b >= a: continue  # Якорь after должен быть после якоря before
                    dist = a - b
                    if best_dist is None or dist < best_dist:
                        best_dist = dist
                        best = (b, a)
            if best:
                b, a = best
                max_doc_distance = 8
                if a - b > max_doc_distance:
                    pos_after = []  # Слишком далеко, сбросить
                else:
                    if (a - b == 1) or next_is_placeholder:
                        pass
                    else:
                        for k in range(b + 1, a):
                            mid = doc_paragraphs[k].strip()
                            if candidate_ok(mid, [anchor_before, anchor_after]):
                                return True, mid, k

        # 2. Поиск по одному якорю (anchor_before)
        if pos_before:
            for b in pos_before:
                para_b = doc_paragraphs[b] or ""
                low_b = para_b.lower()
                ab = anchor_before.strip().lower()
                idx_b = low_b.find(ab)
                if idx_b != -1:
                    after_b = para_b[idx_b + len(ab):].strip()
                    if candidate_ok(after_b, [anchor_before, anchor_after]):
                        return True, after_b, b

                max_forward_search = 10
                for k in range(b + 1, min(len(doc_paragraphs), b + 1 + max_forward_search)):
                    cand_raw = doc_paragraphs[k]

                    if cand_raw is None or cand_raw.strip() == "":
                        continue

                    cand = cand_raw.strip()
                    cand_lower = cand.lower()

                    if any(sw in cand_lower for sw in ["введение", "заключение", "список", "приложение"]) or re.search(
                            r"^\(?\s*(подпись|иниц|инициалы|фамил)", cand_lower):
                        break

                    if candidate_ok(cand, [anchor_before, anchor_after]):
                        return True, cand, k

                    break

        # 3. Поиск по одному якорю (anchor_after)
        if pos_after:
            for a in pos_after:
                para_a = doc_paragraphs[a] or ""
                low_a = para_a.lower()
                aa = anchor_after.strip().lower()
                idx_a = low_a.find(aa)
                if idx_a != -1:
                    before_a = para_a[:idx_a].strip()
                    if candidate_ok(before_a, [anchor_before, anchor_after]):
                        return True, before_a, a

                k = a - 1
                if k >= 0:
                    cand_raw = doc_paragraphs[k]
                    if cand_raw and cand_raw.strip():
                        cand = cand_raw.strip()
                        if any(sw in cand.lower() for sw in stop_words_list) or re.search(
                                r"^\(?\s*(подпись|иниц|инициалы|фамил)", cand.lower()):
                            continue
                        if candidate_ok(cand, [anchor_before, anchor_after]):
                            return True, cand, k

        return False, None, -1

    # Оценка произвольного условия
    def _evaluate_group_condition(self, condition: str, sum_val: float, num_values: int) -> tuple:
        """Выполняет проверку группового условия (SUM, AVG) с операторами."""
        if not condition:
            return True, "Нет условия для проверки."

        condition = condition.strip().upper().replace(' ', '')

        # Шаблон: (SUM|AVG)([<=>!]+)(\d+(\.\d+)?)
        m = re.match(r"(SUM|AVG)([<=>!]+)(\d+(\.\d+)?)", condition)

        if not m:
            return False, f"Условие '{condition}' не распознано или не поддерживается."

        check_type = m.group(1)
        operator = m.group(2)
        target = float(m.group(3))

        value_to_check = sum_val
        check_name = "Сумма"

        if check_type == "AVG":
            if num_values == 0:
                return False, "Невозможно вычислить среднее: в группе нет валидных чисел."
            value_to_check = sum_val / num_values
            check_name = "Среднее"

        # Выполнение сравнения с учетом допуска (tolerance)
        tolerance = 0.001
        is_valid = False

        if operator == '=':
            is_valid = abs(value_to_check - target) < tolerance
        elif operator == '>=':
            is_valid = value_to_check >= target
        elif operator == '<=':
            is_valid = value_to_check <= target
        elif operator == '>':
            is_valid = value_to_check > target
        elif operator == '<':
            is_valid = value_to_check < target
        elif operator == '!=':
            is_valid = abs(value_to_check - target) >= tolerance

        result_val_str = f"{value_to_check:.2f}"

        if is_valid:
            message = f"{check_name} ({result_val_str}) соответствует условию {check_type}{operator}{target}."
        else:
            message = f"{check_name} ({result_val_str}) не соответствует условию {check_type}{operator}{target}."

        return is_valid, message

    # МЕТОД ДЛЯ ГРУПП
    def _check_groups(self, results: list) -> list:
        """Проверяет групповые условия, используя произвольные условия из шаблона."""

        groups_to_check = {}

        for r in results:
            group_name = r.get("group_name")
            group_condition = r.get("group_condition", "").strip()

            if group_name and group_condition:
                # Ключ включает условие, чтобы группы с разными условиями обрабатывались отдельно
                key = (group_name, group_condition)

                if key not in groups_to_check:
                    groups_to_check[key] = {
                        "condition": group_condition,
                        "valid_values": [],
                        "total_sum": 0.0,
                        "all_ok": True,
                        "missing_fields": []
                    }

                # Проверяем, что поле найдено и имеет числовой тип
                if r['status'] == 'ok' and r['expected_type'] == 'number':
                    try:
                        val_str = str(r['value']).replace(' ', '').replace(',', '.')
                        float_value = float(val_str)
                        groups_to_check[key]["valid_values"].append(float_value)
                        groups_to_check[key]["total_sum"] += float_value
                    except ValueError:
                        groups_to_check[key]["all_ok"] = False
                        groups_to_check[key]["missing_fields"].append(f"{r['field']} (не число)")
                else:
                    # Если поле не найдено, невалидно или не числовое — группа не может быть проверена
                    groups_to_check[key]["all_ok"] = False
                    groups_to_check[key]["missing_fields"].append(r['field'])

        # 2. Проверяем каждую группу
        group_report = []

        for (group_name, condition), data in groups_to_check.items():
            sum_val = data["total_sum"]
            num_values = len(data["valid_values"])

            if data["all_ok"]:
                # Если все поля найдены и валидны, оцениваем заданное условие
                is_valid, message = self._evaluate_group_condition(
                    condition, sum_val, num_values
                )

                status = "group_ok" if is_valid else "group_condition_invalid"
            else:
                status = "group_check_failed"
                missing_list = ", ".join(data['missing_fields'])
                message = f"Группа '{group_name}' не проверена: не все поля найдены/валидны. Проблема с полями: {missing_list}."

            group_report.append({
                "field": f"Группа: {group_name}",
                "status": status,
                "value": f"{sum_val:.2f}",
                "group_name": group_name,
                "group_condition": condition,
                "message": message
            })

        return group_report

    def check_document(self, doc_paragraphs: list) -> list:
        results = []
        cursor = 0
        max_len = len(doc_paragraphs)

        # Пропускаем начальные пустые абзацы
        while cursor < max_len and not doc_paragraphs[cursor].strip():
            cursor += 1

        for ph in self.template.get_placeholders():
            name = ph["name"]
            expected_type = ph["type"]
            optional = ph["optional"]
            anchor_before = ph.get("anchor_before", "").strip()
            anchor_after = ph.get("anchor_after", "").strip()
            next_is_placeholder = ph.get("next_is_placeholder", False)
            group_name = ph.get("group_name", "")
            group_condition = ph.get("group_condition", "")

            found, value, found_idx = self._find_value_using_anchors(
                anchor_before, anchor_after, doc_paragraphs, cursor, expected_type, next_is_placeholder
            )

            if not found:
                status = "missing_optional" if optional else "missing"
                results.append(
                    {
                        "field": name,
                        "status": status,
                        "optional": optional,
                        "anchor_before": anchor_before,
                        "anchor_after": anchor_after,
                        "group_name": group_name,
                        "group_condition": group_condition,
                    }
                )
                # Эвристика для табличных данных: если поле пропущено, сдвигаем курсор на 1,
                # чтобы следующее поле в таблице не начало поиск с того же места.
                if expected_type == 'number' and name.startswith("Этап_"):
                    cursor = max(cursor, cursor + 1)
                continue

            # Если найдено, то для числовых полей value - это уже извлеченное число
            is_valid = self._validate_type(value, expected_type) if expected_type != 'number' else True

            results.append(
                {
                    "field": name,
                    "value": value,
                    "expected_type": expected_type,
                    "status": "ok" if is_valid else "invalid",
                    "optional": optional,
                    "anchor_before": anchor_before,
                    "anchor_after": anchor_after,
                    "found_paragraph_index": found_idx,
                    "group_name": group_name,
                    "group_condition": group_condition,
                }
            )

            # Обновляем курсор, чтобы поля шли по порядку.
            cursor = max(cursor, found_idx + 1)
            while cursor < max_len and not doc_paragraphs[cursor].strip():
                cursor += 1

        # Выполнение групповой проверки
        group_results = self._check_groups(results)

        return results + group_results

    # Генерирование отчета
    def generate_report(self, file_name: str, results: list) -> str:
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        header = (
            f"Отчёт проверки\n"
            f"Файл: {file_name}\n"
            f"Шаблон: {os.path.basename(self.template.source_path)}\n"
            f"Дата: {now}\n\n"
        )

        lines = [header, "=== Проверка заполнителей ===\n"]

        group_reports = []

        for r in results:
            if r["field"].startswith("Группа:"):
                group_reports.append(r)
                continue

            group_info = ""
            if r.get('group_name'):
                group_info = f" [Группа: **{r['group_name']}** / Условие: `{r.get('group_condition', '—')}`]"

            if r["status"] == "ok":
                lines.append(f"✅ {r['field']} — найдено: **{r['value']}**{group_info}")
            elif r["status"] == "invalid":
                lines.append(
                    f"⚠️ {r['field']} — найдено, но тип не соответствует ({r['expected_type']}): **{r['value']}**{group_info}"
                )
            elif r["status"] == "missing_optional":
                lines.append(f"ℹ️ {r['field']} — отсутствует (необязательное){group_info}")
            elif r["status"] == "missing":
                lines.append(f"❌ {r['field']} — отсутствует или не найдено корректное значение{group_info}")

        # Добавляем отчеты по группам в конец
        if group_reports:
            lines.append("\n=== Проверка групповых условий ===\n")
            for gr in group_reports:
                if gr['status'] == 'group_ok':
                    lines.append(f"✅ **{gr['field']}** — Успех. {gr['message']}")
                elif gr['status'] == 'group_condition_invalid':
                    lines.append(f"❌ **{gr['field']}** — Ошибка. {gr['message']}")
                elif gr['status'] == 'group_check_failed':
                    lines.append(f"⚠️ **{gr['field']}** — Невозможно проверить. {gr['message']}")

        return "\n".join(lines)


# ============================================================
#  GUI
# ============================================================
class AppGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Проверка документа по шаблону")
        self.root.geometry("780x600")

        self.template = None
        self.document_path = None
        self.document_paragraphs = None
        self.checker = None
        self._last_report_text = ""

        self._build_ui()

    def _build_ui(self):
        frame = tk.Frame(self.root, padx=10, pady=10)
        frame.pack(fill="both", expand=True)

        tk.Label(frame, text="Проверка документа по шаблону [[...]]", font=("Segoe UI", 14, "bold")).pack(
            pady=(0, 8)
        )

        btn_frame = tk.Frame(frame)
        btn_frame.pack(fill="x", pady=6)

        self.load_template_btn = tk.Button(btn_frame, text="Загрузить шаблон", width=18, command=self.load_template)
        self.load_template_btn.pack(side="left", padx=(0, 8))

        self.load_doc_btn = tk.Button(btn_frame, text="Загрузить документ", width=22, state="disabled",
                                      command=self.load_document)
        self.load_doc_btn.pack(side="left", padx=(0, 8))

        self.run_check_btn = tk.Button(btn_frame, text="Проверить", width=12, state="disabled", command=self.run_check)
        self.run_check_btn.pack(side="left")

        self.save_report_btn = tk.Button(btn_frame, text="Сохранить отчёт", width=16, state="disabled",
                                         command=self.save_report)
        self.save_report_btn.pack(side="right")

        self.template_label = tk.Label(frame, text="Шаблон: (не загружен)", anchor="w")
        self.template_label.pack(fill="x")
        self.document_label = tk.Label(frame, text="Документ: (не загружен)", anchor="w")
        self.document_label.pack(fill="x")

        self.result_text = scrolledtext.ScrolledText(frame, wrap=tk.WORD, width=110, height=30, font=("Segoe UI", 10))
        self.result_text.pack(fill="both", expand=True)
        self.result_text.insert(tk.END,
                                "Программа для автоматической проверки самостоятельных работ.\nПожалуйста, загрузите шаблон и документ.")

    def load_template(self):
        path = filedialog.askopenfilename(filetypes=[("Word", "*.docx"), ("PDF", "*.pdf")])
        if not path:
            return
        try:
            tpl = Template.load_from_file(path)
            self.template = tpl
            info = []
            for i, p in enumerate(tpl.get_placeholders(), 1):
                group_name = p['group_name'] if p['group_name'] else '—'
                group_cond = p['group_condition'] if p['group_condition'] else '—'

                info.append(f"{i}. {p['name']} ({p['type']})")

                details_parts = []
                # Добавляем пометку только если optional == True
                if p['optional']:
                    details_parts.append("optional: True")

                details_parts.append(f"group: {group_name}")
                details_parts.append(f"condition: {group_cond}")

                # Собираем строку через запятую
                info.append(f"   {', '.join(details_parts)}")

                info.append(f"   anchor_before: '{p['anchor_before']}'")
                info.append(f"   anchor_after : '{p['anchor_after']}'")
                info.append(f"   para_index: {p['para_index']}\n")

            self.template_label.config(text=f"Шаблон: {os.path.basename(path)} — {len(tpl.get_placeholders())} полей")
            self.result_text.delete(1.0, tk.END)
            self.result_text.insert(tk.END, "\n".join(info))
            self.load_doc_btn.config(state="normal")
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))

    def load_document(self):
        path = filedialog.askopenfilename(filetypes=[("Word", "*.docx"), ("PDF", "*.pdf")])
        if not path:
            return
        try:
            self.document_paragraphs = DocumentLoader.get_paragraphs(path)
            self.document_path = path
            self.document_label.config(text=f"Документ: {os.path.basename(path)}")
            self.run_check_btn.config(state="normal")
        except Exception as e:
            messagebox.showerror("Ошибка загрузки", str(e))

    def run_check(self):
        if not self.template or not self.document_paragraphs:
            messagebox.showwarning("Ошибка", "Сначала загрузите шаблон и документ.")
            return
        try:
            self.checker = DocumentChecker(self.template)
            results = self.checker.check_document(self.document_paragraphs)
            report = self.checker.generate_report(os.path.basename(self.document_path), results)
            self.result_text.delete(1.0, tk.END)
            self.result_text.insert(tk.END, report)
            self._last_report_text = report
            self.save_report_btn.config(state="normal")
        except Exception as e:
            messagebox.showerror("Ошибка проверки", f"Произошла ошибка при выполнении проверки: {e}")

    def save_report(self):
        if not self._last_report_text:
            return
        path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text", "*.txt")])
        if not path:
            return
        with open(path, "w", encoding="utf-8") as f:
            f.write(self._last_report_text)
        messagebox.showinfo("Готово", f"Отчёт сохранён: {os.path.basename(path)}")


# ============================================================
#  main
# ============================================================
def main():
    root = tk.Tk()
    app = AppGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()


